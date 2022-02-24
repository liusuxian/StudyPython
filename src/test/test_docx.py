# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.oxml.ns import qn
from StudyPython.src import docx_utils

# 创建 docx 文件
isExists = os.path.exists('demo.docx')
if isExists:
    os.remove('demo.docx')
# 创建文档对象
document = Document()
# 设置一个空白样式
style = document.styles['Normal']
# 设置西文字体
style.font.name = 'Times New Roman'
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
# 设置文档标题，中文要用unicode字符串
document.add_heading(u'招标公告信息', 0)

p = document.add_paragraph('无锡华侨城梁溪区XDG-2017-27号地块三期景观绿化工程招标公告\n')
p.add_run('北京市\n').bold = True
# p.add_run('https://m.bidcenter.com.cn/newsrili-1-165684150.html\n').italic = True
# 添加链接到url
docx_utils.addHyperlink(p, 'https://m.bidcenter.com.cn/newsrili-1-165684150.html',
                        'https://m.bidcenter.com.cn/newsrili-1-165684150.html', None, True)

p = document.add_paragraph('无锡华侨城梁溪区XDG-2017-27号地块三期景观绿化工程招标公告\n')
p.add_run('北京市\n').bold = True
# p.add_run('https://m.bidcenter.com.cn/newsrili-1-165684150.html\n').italic = True
# 添加链接到url
docx_utils.addHyperlink(p, 'https://m.bidcenter.com.cn/newsrili-1-165684150.html',
                        'https://m.bidcenter.com.cn/newsrili-1-165684150.html', None, True)

document.save('demo.docx')
