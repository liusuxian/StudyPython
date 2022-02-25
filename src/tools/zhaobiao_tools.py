# -*- coding: utf-8 -*-
import calendar
import datetime
import os
import time
import re
from docx import Document
from docx.oxml.ns import qn
from StudyPython.src import docx_utils
import requests

cityList = [
    '成都', '四川',
    '潍坊', '山东',
    '唐山', '河北',
    '西安', '陕西',
    '南京', '江苏'
]


def handelDateUrl(date: str):
    dateUrl = "https://www.bidcenter.com.cn/newsmore-" + date + '.html'
    print('dateUrl:', dateUrl)
    time.sleep(5)
    try:
        result = requests.get(dateUrl)
    except Exception as e:
        print('handelDateUrl ERROR:', e)
    else:
        if result.status_code == 200:
            resUrlList = re.findall('<li><a href="(.*?)".*?>.*?</a>', result.text, re.S)
            for resUrl in resUrlList:
                handelResUrl('https://www.bidcenter.com.cn' + resUrl, date)
        else:
            print('handelDateUrl ERROR:', result.status_code, dateUrl)


def handelResUrl(resUrl: str, date: str):
    print('handelResUrl:', resUrl)
    time.sleep(5)
    try:
        result = requests.get(resUrl)
    except Exception as e:
        print('handelResUrl ERROR:', e)
    else:
        if result.status_code == 200:
            realUrl = re.findall('<link rel.*?href="(.*?)" />', result.text, re.S)
            if len(realUrl) > 0 and 'http' in realUrl[0]:
                handelRealUrl(realUrl[0], date)
        else:
            print('handelResUrl ERROR:', result.status_code, resUrl)


def handelRealUrl(realUrl: str, date: str):
    print('handelRealUrl:', realUrl)
    time.sleep(5)
    try:
        result = requests.get(realUrl)
    except Exception as e:
        print('handelRealUrl ERROR:', e)
    else:
        if result.status_code == 200:
            citys = re.findall('<ul class="xiangm-xx">.*?<a(.*?)</li>', result.text, re.S)
            if len(citys) > 0:
                citys = citys[0]
                city = re.findall('>(.*?)</a>', citys, re.S)
                if len(city) > 0:
                    city = str('-').join(city)
                    for c in cityList:
                        if c in city:
                            title = re.findall('<p class="text-title">(.*?)</p>', result.text, re.S)
                            if len(title) > 0:
                                title = title[0].replace("\r\n", "").replace(" ", "")
                                print('time:', date, 'title:', title, 'city:', city, 'realUrl:', realUrl)
                                p = document.add_paragraph(date + ' ' + title + '\n')
                                p.add_run(city + '\n').bold = True
                                docx_utils.addHyperlink(p, realUrl, realUrl, None, True)
                                # 保存文档
                                document.save('zhaobiao.docx')
                                break
        else:
            print('handelRealUrl ERROR:', result.status_code, realUrl)


# 创建 docx 文件
isExists = os.path.exists('zhaobiao.docx')
if isExists:
    os.remove('zhaobiao.docx')
# 创建文档对象
document = Document()
# 设置一个空白样式
style = document.styles['Normal']
# 设置西文字体
style.font.name = 'Times New Roman'
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
# 设置文档标题，中文要用unicode字符串
document.add_heading(u'招标公告信息', 0)

# 爬取最近6个月的数据
today = datetime.date.today()
endYear = today.year
endMonth = today.month
endDay = today.day
print('today:', today)
curYear = endYear
curMonth = endMonth
for i in range(0, 6):
    print('curYear:', curYear, 'curMonth:', curMonth)
    if curMonth == endMonth:
        for day in range(1, endDay + 1):
            handelDateUrl(str(curYear) + '-' + str(curMonth) + '-' + str(day))
    else:
        _, monthRange = calendar.monthrange(curYear, curMonth)
        for day in range(1, monthRange + 1):
            handelDateUrl(str(curYear) + '-' + str(curMonth) + '-' + str(day))

    curMonth = curMonth - 1
    if curMonth <= 0:
        curYear = curYear - 1
        curMonth = 12
