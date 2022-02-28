# -*- coding: utf-8 -*-
import calendar
import datetime
import os
import time
import re
from docx import Document
from docx.oxml.ns import qn
from StudyPython.src import docx_utils
import requests

cityList = [
    '四川', '成都',
    '山东', '潍坊',
    '河北', '唐山',
    '陕西', '西安',
    '江苏', '南京',
]


def handelDateUrl(date: str):
    time.sleep(1)
    dateUrl = "https://www.bidcenter.com.cn/newsmore-" + date + '.html'
    print('dateUrl:', dateUrl)
    try:
        result = requests.get(dateUrl)
    except Exception as e:
        print('handelDateUrl ERROR:', e)
    else:
        if result.status_code == 200:
            resUrlList = re.findall('<li><a href="(.*?)".*?>.*?</a>', result.text, re.S)
            for resUrl in resUrlList:
                handelUrl('https://www.bidcenter.com.cn' + resUrl, date)
        else:
            print('handelDateUrl ERROR:', result.status_code, dateUrl)


def handelUrl(url: str, date: str):
    time.sleep(1)
    try:
        result = requests.get(url)
    except Exception as e:
        print('handelResUrl ERROR:', e)
    else:
        if result.status_code == 200:
            title = re.findall('<h1 class="jq_lijichakan".*?id="title">(.*?)</h1>', result.text, re.S)
            if len(title) > 0:
                title = title[0].replace("\r\n", "").replace(" ", "")
                print('== time:', date, 'title:', title)
                if ('更新' in title) or ('改造' in title):
                    print('=== time:', date, 'title:', title, 'url:', url)
                    if ('老旧小区' in title) or ('小区' in title) or ('公寓' in title) or ('住区' in title) or ('旧城' in title) or \
                            ('城市' in title):
                        citys = re.findall('<tbody class="zbrl_tb">.*?<a(.*?)</span>', result.text, re.S)
                        if len(citys) > 0:
                            citys = citys[0]
                            city = re.findall('>(.*?)</a>', citys, re.S)
                            print('==== time:', date, 'title:', title, 'city', city, 'url:', url)
                            if len(city) > 0:
                                city = str('-').join(city)
                                for c in cityList:
                                    if c in city:
                                        saveDocx(document, date, title, city, url)
                                        break
        else:
            print('handelUrl ERROR:', result.status_code, url)


# 保存docx文档
def saveDocx(doc, date, title, city, url):
    print('**** time:', date, 'title:', title, 'city:', city, 'url:', url)
    p = doc.add_paragraph(date + ' ' + title + '\n')
    p.add_run(city + '\n').bold = True
    docx_utils.addHyperlink(p, url, url, None, True)
    # 保存文档
    doc.save('zhaobiao.docx')


# 创建 docx 文件
isExists = os.path.exists('zhaobiao.docx')
if isExists:
    os.remove('zhaobiao.docx')
# 创建文档对象
document = Document()
# 设置一个空白样式
style = document.styles['Normal']
# 设置西文字体
style.font.name = 'Times New Roman'
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
# 设置文档标题，中文要用unicode字符串
document.add_heading(u'招标公告信息', 0)

# 爬取最近6个月的数据
today = datetime.date.today()
endYear = today.year
endMonth = today.month
endDay = today.day
print('today:', today)
curYear = endYear
curMonth = endMonth
for i in range(0, 6):
    print('curYear:', curYear, 'curMonth:', curMonth)
    if curMonth == endMonth:
        curDay = endDay
        while curDay > 0:
            handelDateUrl(str(curYear) + '-' + str(curMonth) + '-' + str(curDay))
            curDay = curDay - 1
    else:
        _, monthRange = calendar.monthrange(curYear, curMonth)
        curDay = monthRange
        while curDay > 0:
            handelDateUrl(str(curYear) + '-' + str(curMonth) + '-' + str(curDay))
            curDay = curDay - 1

    curMonth = curMonth - 1
    if curMonth <= 0:
        curYear = curYear - 1
        curMonth = 12
