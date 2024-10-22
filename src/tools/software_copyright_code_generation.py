'''
Author: liusuxian 382185882@qq.com
Date: 2024-10-11 19:36:26
LastEditors: liusuxian 382185882@qq.com
LastEditTime: 2024-10-22 17:28:31
Description: 

Copyright (c) 2024 by liusuxian email: 382185882@qq.com, All Rights Reserved.
'''
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_SECTION_START
from docx import oxml


def create_document(lines, filename):
    docx = Document()
    # 设置字体
    docx.styles['Normal'].font.name = u'宋体'
    docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    docx.styles['Normal'].font.size = Pt(8.5)
    docx.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    # 分节-内容页
    # section = docx.add_section(start_type=WD_SECTION_START.NEW_PAGE)
    section = docx.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    section.is_linked_to_previous = False
    paragraphs = docx.add_paragraph('')

    for line in lines:
        paragraphs.add_run(line)

    # 设置页眉
    header = section.header
    header.paragraphs[0].text = system_name + " " + system_version
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置页脚
    footer = section.footer
    paragraph = footer.paragraphs[0]
    # paragraph.add_run('')
    # 创建页码对象
    page_number = oxml.shared.OxmlElement("w:fldSimple")
    # page_number.set(qn("w:instr"), r'PAGE')
    page_number.set(qn("w:instr"), r'PAGE \* MERGEFORMAT')
    # 创建域对象
    run = oxml.shared.OxmlElement("w:r")
    run.append(page_number)
    # 将域对象插入到段落中
    paragraph._p.append(run)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    docx.save(filename)


# 代码所在文件夹
path = '/Users/liusuxian/Downloads/fasthttp-master'
# 保存路径
save_path = '/Users/liusuxian/Downloads/软著材料/阳光短剧抖音小程序'
# 系统名称
system_name = '阳光短剧抖音小程序'
# 版本号
system_version = 'V1.0'
# 每页的行数
lines_per_page = 60
# 所需的页数
pages_required = 30
# 所需的最大行数
max_lines = lines_per_page * pages_required

all_lines = []
file_lines = []
for root, dirs, files in os.walk(path):
    for file in files:
        if file.endswith('.go') or file.endswith('.html') or file.endswith('.js') or file.endswith('.py'):
            with open(os.path.join(root, file), 'r', encoding='utf-8') as f:
                for line in f:
                    if line.strip() == '':
                        continue
                    file_lines.append(line)
                    if len(file_lines) >= max_lines:
                        all_lines.append(file_lines[:])
                        file_lines = []

if len(file_lines) >= lines_per_page * pages_required:
    all_lines.append(file_lines[:max_lines])

if len(all_lines) > 0:
    create_document(
        all_lines[0], f'{save_path}/{system_name} {system_version} 前30页.docx')
if len(all_lines) > 1:
    create_document(
        all_lines[-1], f'{save_path}/{system_name} {system_version} 后30页.docx')

print('文档生成成功')
